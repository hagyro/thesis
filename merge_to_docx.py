#!/usr/bin/env python3
"""
Script to merge markdown files into existing thesis.docx
Converts markdown formatting to Word styles and inserts at appropriate locations.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys


def add_heading(doc, text, level):
    """Add a heading with proper Greek font support"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
        run.font.bold = True
        # Ensure Greek font compatibility
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return heading


def add_paragraph(doc, text, bold=False, italic=False, style='Normal'):
    """Add a paragraph with proper formatting and Greek support"""
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Parse inline markdown formatting
    parts = parse_inline_markdown(text)

    for part_text, part_bold, part_italic in parts:
        run = para.add_run(part_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold or part_bold
        run.font.italic = italic or part_italic
        # Ensure Greek font compatibility
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    return para


def parse_inline_markdown(text):
    """Parse inline markdown (bold, italic) and return list of (text, bold, italic) tuples"""
    parts = []
    current_pos = 0

    # Pattern for **bold** and *italic*
    pattern = r'(\*\*.*?\*\*|\*.*?\*)'

    for match in re.finditer(pattern, text):
        # Add text before match
        if match.start() > current_pos:
            parts.append((text[current_pos:match.start()], False, False))

        matched_text = match.group(1)
        if matched_text.startswith('**') and matched_text.endswith('**'):
            # Bold text
            parts.append((matched_text[2:-2], True, False))
        elif matched_text.startswith('*') and matched_text.endswith('*'):
            # Italic text
            parts.append((matched_text[1:-1], False, True))

        current_pos = match.end()

    # Add remaining text
    if current_pos < len(text):
        parts.append((text[current_pos:], False, False))

    return parts if parts else [(text, False, False)]


def add_table_from_markdown(doc, lines, start_idx):
    """Parse markdown table and add to document"""
    table_lines = []
    idx = start_idx

    # Collect table lines
    while idx < len(lines) and '|' in lines[idx]:
        if not lines[idx].strip().startswith('|---'):  # Skip separator line
            table_lines.append(lines[idx])
        idx += 1

    if not table_lines:
        return idx

    # Parse table structure
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]  # Remove empty first/last
        rows.append(cells)

    # Create Word table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = 'Light Grid Accent 1'
    except:
        try:
            table.style = 'Table Grid'
        except:
            pass  # Use default style

    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            # Format header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

    doc.add_paragraph()  # Add space after table
    return idx


def process_markdown_file(doc, filepath):
    """Process a markdown file and add to document"""
    print(f"Processing {filepath}...")

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)

        # Tables
        elif '|' in line and i + 1 < len(lines) and '|---' in lines[i + 1]:
            i = add_table_from_markdown(doc, lines, i)
            continue

        # Regular paragraphs
        else:
            if line.strip():
                add_paragraph(doc, line)

        i += 1


def main():
    """Main function to merge markdown files into thesis"""

    # Input/output files
    input_docx = "KΑΡΑΤΖΑ Παρασκευη-.ΔΗΜΙΟΥΡΓΙΑ ΔΙΠΛΩΜΑΤΙΚΗΣ.docx"
    output_docx = "KΑΡΑΤΖΑ_ΔΙΠΛΩΜΑΤΙΚΗ_ΠΛΗΡΗΣ.docx"

    # Files to merge (in order)
    markdown_files = [
        "ΚΕΦΑΛΑΙΟ_5_ΠΛΗΡΕΣ.md",
        "ΣΥΜΠΕΡΑΣΜΑΤΑ_ΠΛΗΡΗ.md",
        "ΒΙΒΛΙΟΓΡΑΦΙΑ.md"
    ]

    print(f"Loading {input_docx}...")
    try:
        doc = Document(input_docx)
    except Exception as e:
        print(f"Error loading document: {e}")
        print("Creating new document instead...")
        doc = Document()

    # Add page break before new content
    doc.add_page_break()

    # Process each markdown file
    for md_file in markdown_files:
        try:
            process_markdown_file(doc, md_file)
            doc.add_page_break()  # Page break between major sections
        except FileNotFoundError:
            print(f"Warning: {md_file} not found, skipping...")
        except Exception as e:
            print(f"Error processing {md_file}: {e}")

    # Save the merged document
    print(f"Saving to {output_docx}...")
    doc.save(output_docx)
    print(f"✓ Successfully created {output_docx}")
    print(f"\nThe complete thesis is now in: {output_docx}")


if __name__ == "__main__":
    main()
